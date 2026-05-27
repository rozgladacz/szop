"""A4.1 — testy parsera DOCX → rules_extracted.yaml.

Pokrycie:
- Slug generation (`make_slug`) — golden cases włącznie z Polish chars i `Ł/ł` (non-decomp).
- Pełen `SZOP.docx` sanity: count in expected range, unique slugs, no empty fields,
  types z dozwolonego zbioru.
- Golden test: programmatic-generated fixture DOCX z 2 zdolności + oczekiwany schema.
- Edge cases: missing file (exit 1), invalid file (exit 1), empty DOCX (empty list).
- Embedded newlines: dwie zdolności w jednym paragrafie split przez `\n`.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest
import yaml
from docx import Document

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from scripts.rules_extract import (  # noqa: E402
    extract_abilities,
    main,
    make_slug,
    validate_uniqueness,
)


SCRIPT_PATH = ROOT_DIR / "scripts" / "rules_extract.py"
REAL_DOCX = ROOT_DIR / "app" / "static" / "docs" / "SZOP.docx"


# --- Slug generation --------------------------------------------------------


@pytest.mark.parametrize(
    "name, expected_slug",
    [
        ("Bohater", "bohater"),
        ("Zasadzka", "zasadzka"),
        ("Szybki/Wolny", "szybki_wolny"),
        ("Mag(X)", "mag"),
        ("Transport(X)", "transport"),
        ("Mistrzostwo(X)", "mistrzostwo"),
        ("Dobrze/źle strzela", "dobrze_zle_strzela"),
        # NFKD edge cases — Polish chars that DO decompose:
        ("Ostrożny", "ostrozny"),
        ("Latający", "latajacy"),
        ("Niewrażliwy", "niewrazliwy"),
        # Non-decomposing chars (Ł/ł must be pre-replaced):
        ("Łatanie", "latanie"),
        ("Ociężałość", "ociezalosc"),
        ("Przełamanie", "przelamanie"),
        ("Otwarty transport(X)", "otwarty_transport"),
        # Edge: special punctuation
        ("Waagh!", "waagh"),
        # Empty edges
        ("AP(X)", "ap"),
    ],
)
def test_make_slug(name: str, expected_slug: str) -> None:
    assert make_slug(name) == expected_slug


# --- Validation -------------------------------------------------------------


def test_validate_uniqueness_passes_on_unique() -> None:
    from scripts.rules_extract import ExtractedAbility

    abilities = [
        ExtractedAbility(slug="a", name="A", type="passive", description="desc"),
        ExtractedAbility(slug="b", name="B", type="active", description="desc"),
    ]
    # Should not raise
    validate_uniqueness(abilities)


def test_validate_uniqueness_raises_on_duplicate() -> None:
    from scripts.rules_extract import ExtractedAbility

    abilities = [
        ExtractedAbility(slug="a", name="A1", type="passive", description="d1"),
        ExtractedAbility(slug="a", name="A2", type="active", description="d2"),
    ]
    with pytest.raises(ValueError, match="Duplicate slugs"):
        validate_uniqueness(abilities)


# --- Real DOCX sanity -------------------------------------------------------


@pytest.fixture(scope="module")
def real_abilities() -> list:
    return extract_abilities(REAL_DOCX)


def test_real_docx_count_in_expected_range(real_abilities: list) -> None:
    """Sanity: znamy ~87 abilities w `ABILITY_DEFINITIONS`; DOCX może odbiegać o ±10
    przez compound names (Szybki/Wolny), różne slugi (burzaca/przelamanie)."""
    n = len(real_abilities)
    assert 80 <= n <= 100, f"Expected 80-100 abilities, got {n} — parser regression?"


def test_real_docx_all_slugs_unique(real_abilities: list) -> None:
    slugs = [a.slug for a in real_abilities]
    assert len(slugs) == len(set(slugs)), "Duplicate slugs in extracted abilities"


def test_real_docx_all_fields_nonempty(real_abilities: list) -> None:
    for ab in real_abilities:
        assert ab.slug, f"Empty slug for {ab.name!r}"
        assert ab.name, f"Empty name for slug={ab.slug}"
        assert ab.description, f"Empty description for {ab.slug}"
        assert len(ab.description) >= 5, f"Suspiciously short desc for {ab.slug}: {ab.description!r}"


def test_real_docx_types_from_valid_set(real_abilities: list) -> None:
    valid_types = {"passive", "active", "aura", "weapon"}
    for ab in real_abilities:
        assert ab.type in valid_types, f"Invalid type {ab.type!r} for {ab.slug}"


def test_real_docx_known_abilities_present(real_abilities: list) -> None:
    """Smoke: kilka stabilnych zdolności musi być w extract."""
    slugs = {a.slug for a in real_abilities}
    must_have = {"bohater", "zasadzka", "zwiadowca", "transport", "mag", "ap"}
    missing = must_have - slugs
    assert not missing, f"Missing core abilities: {missing}"


def test_real_docx_type_distribution(real_abilities: list) -> None:
    """Sanity: oczekujemy że passive dominuje (~60%), weapon ~25%, active+aura reszta."""
    by_type: dict[str, int] = {}
    for ab in real_abilities:
        by_type[ab.type] = by_type.get(ab.type, 0) + 1
    assert by_type.get("passive", 0) >= 30, f"Suspicious type distribution: {by_type}"
    assert by_type.get("weapon", 0) >= 15, f"Suspicious type distribution: {by_type}"
    assert by_type.get("active", 0) >= 5, f"Suspicious type distribution: {by_type}"
    assert by_type.get("aura", 0) >= 2, f"Suspicious type distribution: {by_type}"


# --- Golden test (programmatic fixture) -------------------------------------


@pytest.fixture
def minimal_docx(tmp_path: Path) -> Path:
    """Stwórz minimalną fixture DOCX z 2 zdolnościami pod sekcją Pasywne:."""
    doc = Document()
    doc.add_paragraph("Ignored game rules section.")
    doc.add_paragraph("More rules to skip.")
    doc.add_paragraph("Pasywne:")
    doc.add_paragraph("Zdolność A: Pierwsza zdolność testowa.")
    doc.add_paragraph("Zdolność B(X): Druga z parametrem X.")
    doc.add_paragraph("Aktywne:")
    doc.add_paragraph("Akcja C: Trzecia jest aktywna.")
    # Embedded newline (Shift+Enter) — two abilities in one paragraph
    p = doc.add_paragraph("Akcja D: Czwarta na początku.")
    p.add_run("\nAkcja E: Piąta po Shift+Enter.")
    doc.add_paragraph("Koszt oddziału jest sumą kosztów modeli w nim.")
    doc.add_paragraph("Should be ignored after stop marker.")

    path = tmp_path / "minimal.docx"
    doc.save(str(path))
    return path


def test_golden_minimal_docx(minimal_docx: Path) -> None:
    abilities = extract_abilities(minimal_docx)

    assert len(abilities) == 5, f"Expected 5 abilities, got {len(abilities)}: {[a.name for a in abilities]}"

    by_name = {a.name: a for a in abilities}
    assert "Zdolność A" in by_name
    assert by_name["Zdolność A"].type == "passive"
    assert by_name["Zdolność A"].slug == "zdolnosc_a"
    assert by_name["Zdolność A"].description == "Pierwsza zdolność testowa."

    assert "Zdolność B(X)" in by_name
    assert by_name["Zdolność B(X)"].slug == "zdolnosc_b"
    assert by_name["Zdolność B(X)"].type == "passive"

    assert "Akcja C" in by_name
    assert by_name["Akcja C"].type == "active"

    # Embedded \n split — both should be present
    assert "Akcja D" in by_name
    assert by_name["Akcja D"].type == "active"
    assert "Akcja E" in by_name
    assert by_name["Akcja E"].type == "active"


def test_golden_cli_smoke(minimal_docx: Path, tmp_path: Path) -> None:
    """Run via CLI entry point — should exit 0 + create file."""
    output = tmp_path / "out.yaml"
    rc = main(["--input", str(minimal_docx), "--output", str(output)])
    assert rc == 0
    assert output.exists()

    with open(output, encoding="utf-8") as fh:
        data = yaml.safe_load(fh)
    assert data["version"] == 1
    assert "source" in data
    assert len(data["abilities"]) == 5


# --- Error handling --------------------------------------------------------


def test_missing_file_returns_exit_1(tmp_path: Path) -> None:
    """Wywołanie CLI z nieistniejącym plikiem → exit 1."""
    result = subprocess.run(
        [sys.executable, str(SCRIPT_PATH), "--input", str(tmp_path / "missing.docx"), "--output", str(tmp_path / "out.yaml")],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 1
    assert "ERROR" in result.stderr
    assert "not found" in result.stderr.lower() or "no such" in result.stderr.lower()


def test_invalid_docx_returns_exit_1(tmp_path: Path) -> None:
    """Wywołanie z plikiem nie-DOCX (txt) → exit 1 z czytelnym message."""
    bad = tmp_path / "notdocx.docx"
    bad.write_text("this is not a docx", encoding="utf-8")
    result = subprocess.run(
        [sys.executable, str(SCRIPT_PATH), "--input", str(bad), "--output", str(tmp_path / "out.yaml")],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 1
    assert "ERROR" in result.stderr


def test_empty_docx_returns_empty_list(tmp_path: Path) -> None:
    """DOCX bez `Pasywne:` header — start marker nigdy nie trafiony, pusta lista."""
    doc = Document()
    doc.add_paragraph("Just some text.")
    doc.add_paragraph("No section headers here.")
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    abilities = extract_abilities(path)
    assert abilities == []
