"""Extract ability definitions from DOCX into a YAML drift snapshot.

Pierwszy etap pipeline'u drift-only (ADR-0006). Czyta paragrafy z `SZOP.docx`,
emituje schema kompatybilny z `app/rulesets/v1/abilities.yaml`. `cost_fn` /
wartości kosztów pozostają poza scope — drift sprawdza tylko shape
(slug/name/type/description), nie liczby.

Schema wynikowy: `RulesetAbility` z `app/services/rulesets/models.py` (reuse,
nie duplikat) — pole `value_*` zawsze None bo DOCX ich nie nosi.

Parser jest content-based (DOCX nie ma Headingów, tylko `Normal`/`List
Paragraph`). Sekcje rozpoznawane przez paragrafy końca dwukropkiem:
`Pasywne:`/`Aktywne:`/`Aury:`/`Broni:`. Word soft line break (Shift+Enter)
emituje `\\n` wewnątrz `paragraph.text` — kilka zdolności potrafi dzielić
jeden paragraf, dlatego rozbijamy `text.split("\\n")`.

Użycie:
    python scripts/rules_extract.py
    python scripts/rules_extract.py --input app/static/docs/SZOP.docx --output build/rules_extracted.yaml
"""

from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from pathlib import Path

import yaml
from docx import Document
from pydantic import BaseModel, ConfigDict, Field

# Pozwala uruchamiać skrypt bezpośrednio (`python scripts/rules_extract.py`) —
# bez tego `import app.*` failuje gdy CWD != project root.
_PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

# Reuse istniejący schema — A4 jest klientem rulesets, nie tworzy własnego.
from app.services.rulesets.models import AbilityType, RulesetAbility  # noqa: E402


class RulesExtract(BaseModel):
    """Korzeń `build/rules_extracted.yaml` — wersja + ścieżka źródła + lista."""

    model_config = ConfigDict(frozen=True, extra="forbid")

    version: int = Field(default=1, ge=1)
    source: str
    abilities: tuple[RulesetAbility, ...]


# --- Parser constants -------------------------------------------------------

SECTION_MAP: dict[str, AbilityType] = {
    "Pasywne": "passive",
    "Aktywne": "active",
    "Aury": "aura",
    "Broni": "weapon",
}

# Pattern: <Name> (no colon/period inside) : <description>
# Name: 2-60 chars; description: ≥1 char.
NAME_DESC = re.compile(r"^([^:.]{2,60}?):\s+(.+)$")

START_MARKER = "Pasywne:"
STOP_MARKER = "Koszt oddziału jest sumą"

# Wielowyrazowe nagłówki sekcji bez własnych zdolności — flush + skip.
IGNORED_HEADERS = {
    "Dodatkowe zdolności:",
    "Zasady Armii: (zdolności pasywne wycenianie przy założeniu, że ma je prawie każdy oddział w rozpisce)",
}


# --- Slug generation --------------------------------------------------------

# Polish `Ł`/`ł` to osobne Latin chars w Unicode, nie precomposed z combining
# marks — NFKD nie rozkłada ich, więc `encode("ascii", "ignore")` je dropuje.
_POLISH_NONDECOMP = {"ł": "l", "Ł": "L"}


def make_slug(name: str) -> str:
    """Deterministic slug: NFKD → strip accents → lowercase → spaces→underscore.

    `(X)` parameter notation jest odrzucana; slashe i spacje stają się
    underscorem. Slug musi być stabilnym identyfikatorem niezależnym od
    formatowania DOCX.

    Examples:
        "Bohater" → "bohater"
        "Szybki/Wolny" → "szybki_wolny"
        "Mag(X)" → "mag"
        "Mistrzostwo(X)" → "mistrzostwo"
        "Dobrze/źle strzela" → "dobrze_zle_strzela"
        "Łatanie" → "latanie"
        "Ociężałość" → "ociezalosc"
    """
    cleaned = re.sub(r"\(\w+\)", "", name).strip()
    for orig, repl in _POLISH_NONDECOMP.items():
        cleaned = cleaned.replace(orig, repl)
    decomposed = unicodedata.normalize("NFKD", cleaned)
    ascii_only = decomposed.encode("ascii", "ignore").decode("ascii")
    lowered = ascii_only.lower()
    underscored = re.sub(r"[/\s]+", "_", lowered)
    return re.sub(r"[^a-z0-9_]", "", underscored).strip("_")


# --- Extractor --------------------------------------------------------------

class _PendingAbility:
    """Mutable accumulator dla bieżącej zdolności podczas parsingu.

    Trzymamy mutable dict zamiast frozen `RulesetAbility` żeby uniknąć
    per-paragraph rebuilds; finalna immutable instancja powstaje raz przy
    flush.
    """

    __slots__ = ("slug", "name", "type", "description_parts")

    def __init__(self, slug: str, name: str, type: AbilityType, description: str) -> None:
        self.slug = slug
        self.name = name
        self.type = type
        self.description_parts: list[str] = [description]

    def append(self, text: str) -> None:
        self.description_parts.append(text)

    def finalize(self) -> RulesetAbility:
        return RulesetAbility(
            slug=self.slug,
            name=self.name,
            type=self.type,
            description=" ".join(self.description_parts).strip(),
        )


def extract_abilities(docx_path: Path) -> list[RulesetAbility]:
    """Walk paragraphs in DOCX, return list of extracted abilities.

    Raises:
        RuntimeError: jeśli DOCX nie da się otworzyć (`python-docx` exception
            opakowywana w czytelny message; sam plik-nie-istnieje przepuszczamy
            `Document(...)` żeby nie duplikować TOCTOU check).
    """

    try:
        document = Document(str(docx_path))
    except FileNotFoundError:
        # Przekazujemy oryginalny FileNotFoundError — CLI obsłuży w main().
        raise
    except Exception as exc:
        raise RuntimeError(f"Failed to open DOCX (not a valid .docx?): {exc}") from exc

    started = False
    current_section: AbilityType | None = None
    abilities: list[RulesetAbility] = []
    pending: _PendingAbility | None = None

    def flush() -> None:
        nonlocal pending
        if pending is not None:
            abilities.append(pending.finalize())
        pending = None

    for paragraph in document.paragraphs:
        # Word soft line break (Shift+Enter) emituje `\n` wewnątrz paragraph.text.
        # Wiele zdolności potrafi dzielić jeden paragraf (np. Porażenie/Zguba/Dezintegracja).
        for raw_line in paragraph.text.split("\n"):
            text = raw_line.strip()
            if not text:
                continue

            # Faza 1: skip wszystkiego przed pierwszym `Pasywne:`.
            if not started:
                if text == START_MARKER:
                    started = True
                    current_section = SECTION_MAP["Pasywne"]
                continue

            # Faza 2: stop na sekcji formuł kosztu.
            if text.startswith(STOP_MARKER):
                flush()
                return abilities

            # Faza 3: nagłówek sekcji single-word — przełącz current_section.
            section_match = re.match(r"^(\w+):$", text)
            if section_match and section_match.group(1) in SECTION_MAP:
                flush()
                current_section = SECTION_MAP[section_match.group(1)]
                continue

            # Faza 4: nagłówek wielowyrazowy do zignorowania (Dodatkowe zdolności / Zasady Armii).
            if text in IGNORED_HEADERS:
                flush()
                continue

            # Faza 5: linia "Name: desc" → nowa zdolność.
            match = NAME_DESC.match(text)
            if match and current_section is not None:
                flush()
                name = match.group(1).strip()
                pending = _PendingAbility(
                    slug=make_slug(name),
                    name=name,
                    type=current_section,
                    description=match.group(2).strip(),
                )
                continue

            # Faza 6: linia continuation — dolep do bieżącej zdolności.
            if pending is not None:
                pending.append(text)

    flush()
    return abilities


# --- Validation -------------------------------------------------------------

def validate_uniqueness(abilities: list[RulesetAbility]) -> None:
    """Raise `ValueError` z listą duplikatów slug → name1 vs name2."""
    seen: dict[str, str] = {}
    duplicates: list[tuple[str, str, str]] = []
    for ab in abilities:
        if ab.slug in seen:
            duplicates.append((ab.slug, seen[ab.slug], ab.name))
        else:
            seen[ab.slug] = ab.name
    if duplicates:
        lines = ["Duplicate slugs detected:"]
        for slug, name1, name2 in duplicates:
            lines.append(f"  {slug}: '{name1}' vs '{name2}'")
        raise ValueError("\n".join(lines))


# --- YAML serializer --------------------------------------------------------

def write_yaml(extract: RulesExtract, output_path: Path) -> None:
    """Serialize do YAML z konwencjami `abilities.yaml` (UTF-8, no sort, wide)."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "version": extract.version,
        "source": extract.source,
        "abilities": [
            {
                "slug": ab.slug,
                "name": ab.name,
                "type": ab.type,
                "description": ab.description,
            }
            for ab in extract.abilities
        ],
    }
    with open(output_path, "w", encoding="utf-8") as fh:
        yaml.safe_dump(
            payload,
            fh,
            allow_unicode=True,
            sort_keys=False,
            width=10000,
        )


# --- CLI --------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Extract ability definitions from SZOP.docx → build/rules_extracted.yaml",
    )
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("app/static/docs/SZOP.docx"),
        help="Path to SZOP.docx (default: app/static/docs/SZOP.docx)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("build/rules_extracted.yaml"),
        help="Output YAML path (default: build/rules_extracted.yaml)",
    )
    args = parser.parse_args(argv)

    try:
        abilities = extract_abilities(args.input)
    except FileNotFoundError as exc:
        print(f"ERROR: DOCX not found: {exc.filename or args.input}", file=sys.stderr)
        return 1
    except RuntimeError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    try:
        validate_uniqueness(abilities)
    except ValueError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    extract = RulesExtract(
        version=1,
        source=str(args.input).replace("\\", "/"),
        abilities=tuple(abilities),
    )

    write_yaml(extract, args.output)
    print(
        f"Extracted {len(abilities)} abilities from {args.input} → {args.output}",
        file=sys.stderr,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
